# fastfilefinder_scan.py - fast content scanner with optional Office document support
# Outputs UTF-8 TSV lines: path \t entry \t lineno \t snippet

import argparse
import os
import platform
import re
import shutil
import subprocess
import sys
import threading
import time
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from io import TextIOWrapper
from zipfile import BadZipFile, ZipFile
from typing import Iterable, List, Optional, Tuple

# Optional dependencies
try:  # python-docx for .docx
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover - dependency may be missing
    Document = None  # type: ignore

try:  # openpyxl for .xlsx
    import openpyxl  # type: ignore
    from openpyxl.utils.cell import get_column_letter  # type: ignore
except Exception:  # pragma: no cover - dependency may be missing
    openpyxl = None  # type: ignore
    get_column_letter = None  # type: ignore

try:  # xlrd for legacy .xls (requires <=1.2)
    import xlrd  # type: ignore

    def _xlrd_supports_xls() -> bool:
        try:
            version = getattr(xlrd, "__version__", "0")
            parts = [int(p) for p in version.split(".")[:2]]
            return len(parts) >= 2 and (parts[0], parts[1]) < (2, 0)
        except Exception:
            return True

    if not _xlrd_supports_xls():  # pragma: no cover - depends on environment
        xlrd = None  # type: ignore
except Exception:  # pragma: no cover - dependency may be missing
    xlrd = None  # type: ignore

try:  # pywin32 for legacy .doc via Word COM (Office と Python のビット数 32/64 を一致させる必要あり)
    import pythoncom  # type: ignore
    import win32com.client  # type: ignore
    import pywintypes  # type: ignore
except Exception:  # pragma: no cover - dependency may be missing
    pythoncom = None  # type: ignore
    win32com = None  # type: ignore
    pywintypes = None  # type: ignore

# Required optional packages: python-docx, openpyxl, "xlrd<2.0", pywin32


# Common encodings to try for text files
ENCODINGS = (
    "utf-8-sig",
    "utf-16-le",
    "utf-16-be",
    "cp932",
    "utf-8",
)

# Ensure stdout is UTF-8
try:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace", line_buffering=True)
except Exception:
    pass

try:
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8", errors="replace", line_buffering=True)
except Exception:
    pass

stdout_lock = threading.Lock()
word_diag_lock = threading.Lock()
warned = set()

DOC_DIAG_CONTEXT = {
    "enabled": False,
    "perfile": 0,
    "exts": "",
    "legacy_mode": "com",
    "single_thread": False,
}
DOC_DIAG_LOCK = threading.Lock()
DOC_DIAG_EMITTED = False

DOC_WORKER_POOL: Optional["DocWorkerPool"] = None


def ensure_extended_path(path: str) -> str:
    """Return a path with the Windows long path prefix when needed."""

    normalized = os.path.abspath(path)
    if os.name != "nt":
        return normalized

    if normalized.startswith("\\\\?\\"):
        return normalized
    if normalized.startswith("\\\\"):
        return "\\\\?\\UNC\\" + normalized[2:]
    if len(normalized) >= 248:
        return "\\\\?\\" + normalized
    return normalized


def _format_hresult(error: Exception) -> str:
    hresult = getattr(error, "hresult", None)
    if hresult is None:
        return ""
    if hresult < 0:
        hresult &= 0xFFFFFFFF
    return f", hresult=0x{hresult:08X}"


def _summarize_open_args(args: Optional[dict]) -> str:
    if not args:
        return ""
    show_keys = [
        "ReadOnly",
        "AddToRecentFiles",
        "ConfirmConversions",
        "Visible",
        "Encoding",
        "Revert",
    ]
    parts = []
    for key in show_keys:
        if key in args:
            parts.append(f"{key}={args[key]}")
    if not parts:
        return ""
    return ", args=" + ", ".join(parts)


class DocConversionError(Exception):
    def __init__(self, stage: str, detail: Optional[str] = None):
        super().__init__(detail or stage)
        self.stage = stage
        self.detail = detail or ""


def _clean_detail(detail: Optional[str]) -> Optional[str]:
    if not detail:
        return None
    cleaned = " ".join(str(detail).split())
    if len(cleaned) > 500:
        cleaned = cleaned[:497] + "..."
    return cleaned


def _log_doc_stage(stage: str, detail: str) -> None:
    message = f"LOG .doc [{stage}] {detail}"
    sys.stderr.write(message + "\n")
    sys.stderr.flush()


def _log_doc_error(stage: str, path: str, detail: Optional[str] = None) -> None:
    detail = _clean_detail(detail)
    message = f"ERR .doc [{stage}] {path}"
    if detail:
        message += f" ({detail})"
    sys.stderr.write(message + "\n")
    sys.stderr.flush()


def _format_com_exception(exc: Exception) -> str:
    parts = []
    hresult = getattr(exc, "hresult", None)
    if isinstance(hresult, int):
        value = hresult if hresult >= 0 else (hresult & 0xFFFFFFFF)
        parts.append(f"hresult=0x{value:08X}")
    strerror = getattr(exc, "strerror", None)
    if strerror:
        parts.append(f"strerror={_clean_detail(strerror)}")
    excepinfo = getattr(exc, "excepinfo", None)
    if excepinfo:
        try:
            cleaned = ",".join(_clean_detail(part) or "" for part in excepinfo if part)
        except Exception:
            cleaned = None
        if cleaned:
            parts.append(f"excepinfo={cleaned}")
    text = _error_text(exc)
    if text:
        parts.append(f"msg={_clean_detail(text)}")
    return ", ".join(parts)


class DocWorkerPool:
    def __init__(self, single_thread: bool):
        self._single_thread = bool(single_thread)
        max_workers = 1 if self._single_thread else max(1, min(4, (os.cpu_count() or 2)))
        self._executor = ThreadPoolExecutor(
            max_workers=max_workers,
            thread_name_prefix="doc-worker",
        )

    def run(self, func, *args, **kwargs):
        def _callable():
            coinitialized = False
            try:
                if pythoncom is not None:
                    try:
                        pythoncom.CoInitialize()
                    except Exception as exc:
                        raise DocConversionError(
                            "Init",
                            _format_com_exception(exc) or str(exc),
                        )
                    coinitialized = True
                return func(*args, **kwargs)
            finally:
                if coinitialized:
                    try:
                        pythoncom.CoUninitialize()
                    except Exception as exc:
                        warn_once(
                            "coinitialize_cleanup",
                            f"CoUninitialize で例外: {_error_text(exc)}{_format_hresult(exc)}",
                        )

        future = self._executor.submit(_callable)
        return future.result()

    def shutdown(self) -> None:
        self._executor.shutdown(wait=True)

    @property
    def single_thread(self) -> bool:
        return self._single_thread


def configure_doc_workers(single_thread: bool) -> None:
    global DOC_WORKER_POOL
    if DOC_WORKER_POOL is not None:
        DOC_WORKER_POOL.shutdown()
    DOC_WORKER_POOL = DocWorkerPool(single_thread)


def shutdown_doc_workers() -> None:
    global DOC_WORKER_POOL
    pool = DOC_WORKER_POOL
    if pool is not None:
        pool.shutdown()
        DOC_WORKER_POOL = None


def emit_tsv(path: str, entry: str, lineno: int, line: str) -> None:
    line = line.replace("\t", " ").rstrip("\r\n")
    with stdout_lock:
        print(f"{path}\t{entry}\t{lineno}\t{line}", flush=True)


def emit_status(tag: str, *parts: object) -> None:
    payload = "\t".join(str(p) for p in parts)
    with stdout_lock:
        print(f"#{tag}\t{payload}", flush=True)


def warn_once(kind: str, message: str) -> None:
    if kind in warned:
        return
    warned.add(kind)
    sys.stderr.write(message + "\n")
    sys.stderr.flush()


def iter_paths(folder: str, recursive: bool, excluded: set):
    if recursive:
        excluded_lower = {name.lower() for name in excluded}
        for root, dirs, files in os.walk(folder):
            dirs[:] = [d for d in dirs if d.lower() not in excluded_lower]
            for name in files:
                yield os.path.join(root, name)
    else:
        try:
            for name in os.listdir(folder):
                p = os.path.join(folder, name)
                if os.path.isfile(p):
                    yield p
        except Exception:
            return


def normalize_ext(ext: str) -> str:
    return ext.lower().lstrip(".")


def should_target(path: str, exts: set) -> bool:
    if not exts:
        return True
    return normalize_ext(os.path.splitext(path)[1]) in exts


def build_matcher(pattern: str, use_regex: bool):
    if use_regex:
        rx = re.compile(pattern, re.IGNORECASE)

        def matcher(line: str):
            return rx.search(line)

        return matcher

    lowered = pattern.lower()

    def matcher(line: str):
        return lowered in line.lower()

    return matcher


def scan_text_file(path: str, matcher, exts: set, perfile: int) -> int:
    if not should_target(path, exts):
        return 0
    hits = 0
    for enc in ENCODINGS:
        try:
            with open(path, "r", encoding=enc, errors="replace") as reader:
                for lineno, line in enumerate(reader, 1):
                    if matcher(line):
                        emit_tsv(path, "", lineno, line)
                        hits += 1
                        if perfile and hits >= perfile:
                            return hits
            break
        except UnicodeDecodeError:
            continue
        except (OSError, IOError):
            return hits
    return hits


def scan_zip(path: str, matcher, exts: set, perfile: int) -> int:
    hits = 0
    try:
        with ZipFile(path) as zf:
            for name in zf.namelist():
                if exts and normalize_ext(os.path.splitext(name)[1]) not in exts:
                    continue
                entry_hits = 0
                try:
                    for enc in ENCODINGS:
                        try:
                            with zf.open(name, "r") as raw, TextIOWrapper(
                                raw, encoding=enc, errors="replace"
                            ) as reader:
                                for lineno, line in enumerate(reader, 1):
                                    if matcher(line):
                                        emit_tsv(path, name, lineno, line)
                                        entry_hits += 1
                                        hits += 1
                                        if perfile and entry_hits >= perfile:
                                            break
                            break
                        except UnicodeDecodeError:
                            continue
                except KeyError:
                    continue
    except BadZipFile:
        pass
    except Exception as exc:
        warn_once(f"zip:{path}", f"ZIP 読み取り失敗: {path} ({exc})")
    return hits


def iter_docx_lines(doc):
    line_no = 0
    for idx, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip()
        line_no += 1
        yield line_no, f"paragraph:{idx}", text
    for t_idx, table in enumerate(doc.tables, 1):
        for r_idx, row in enumerate(table.rows, 1):
            cells = [cell.text.strip() for cell in row.cells]
            text = "\t".join(cells).strip()
            line_no += 1
            yield line_no, f"table{t_idx}:row{r_idx}", text


def scan_docx(path: str, matcher, perfile: int) -> int:
    if Document is None:
        warn_once("docx", "python-docx がインストールされていないため .docx をスキップします")
        return 0
    hits = 0
    try:
        doc = Document(path)
    except Exception as exc:
        warn_once(f"docx:{path}", f".docx 読み込み失敗: {path} ({exc})")
        return 0
    for lineno, entry, text in iter_docx_lines(doc):
        if not text:
            continue
        if matcher(text):
            emit_tsv(path, entry, lineno, text)
            hits += 1
            if perfile and hits >= perfile:
                break
    return hits


def iter_xlsx_cells(workbook):
    for sheet in workbook.worksheets:
        title = sheet.title
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
            for col_idx, value in enumerate(row, 1):
                if value is None:
                    continue
                text = str(value).strip()
                if not text:
                    continue
                addr = f"{title}!{get_column_letter(col_idx)}{row_idx}" if get_column_letter else f"{title}!{col_idx},{row_idx}"
                yield row_idx, addr, text


def scan_xlsx(path: str, matcher, perfile: int) -> int:
    if openpyxl is None:
        warn_once("xlsx", "openpyxl がインストールされていないため .xlsx をスキップします")
        return 0
    hits = 0
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception as exc:
        warn_once(f"xlsx:{path}", f".xlsx 読み込み失敗: {path} ({exc})")
        return 0
    try:
        for lineno, entry, text in iter_xlsx_cells(wb):
            if matcher(text):
                emit_tsv(path, entry, lineno, text)
                hits += 1
                if perfile and hits >= perfile:
                    break
    finally:
        try:
            wb.close()
        except Exception:
            pass
    return hits


WD_FORMAT_UNICODE_TEXT = 7
MSO_AUTOMATION_SECURITY_FORCE_DISABLE = 3
DOC_LOCK_RETRY_COUNT = 2
DOC_LOCK_RETRY_DELAY = 0.6
DOC_SAVEAS_POLL_INTERVAL = 0.1
DOC_SAVEAS_MAX_WAIT = 10.0

WORD_DIAGNOSTICS_EMITTED = False


def _emit_startup_diag(perfile: int, exts_argument: str, legacy_mode: str, single_thread: bool) -> None:
    arch = platform.architecture()[0]
    if "64" in arch:
        py_bits = "64"
    elif "32" in arch:
        py_bits = "32"
    else:
        py_bits = arch

    word_detect = "NG"
    if pythoncom is not None and win32com is not None:
        try:
            pythoncom.CoInitialize()
        except Exception:
            pass
        else:
            try:
                word = win32com.client.gencache.EnsureDispatch("Word.Application")
            except Exception:
                word_detect = "NG"
            else:
                word_detect = "OK"
                try:
                    word.Visible = False
                except Exception:
                    pass
                try:
                    word.Quit()
                except Exception:
                    pass
            finally:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass

    exts_display = (exts_argument or "").strip() or "*"
    legacy_display = (legacy_mode or "com").lower()
    message = (
        f"diag: py={py_bits}, word-detect={word_detect}, "
        f"perfile={perfile}, exts={exts_display}, "
        f"legacy-doc={legacy_display}, single-thread={bool(single_thread)}"
    )
    sys.stderr.write(message + "\n")
    sys.stderr.flush()


def _error_text(exc: Exception) -> str:
    if pywintypes is not None and isinstance(exc, pywintypes.com_error):
        try:
            if len(exc.args) >= 2 and isinstance(exc.args[1], str) and exc.args[1]:
                return exc.args[1]
        except Exception:
            pass
    return str(exc)


def _looks_like_eula_block(exc: Exception) -> bool:
    text = _error_text(exc).lower()
    return any(keyword in text for keyword in ["eula", "license", "ダイアログ", "dialog"])


def _should_retry_lock(exc: Exception) -> bool:
    hresult = getattr(exc, "hresult", None)
    lock_codes = {
        -2147024864,  # 0x80070020 sharing violation
        -2146823117,  # Word specific sharing violation
    }
    if hresult in lock_codes:
        return True
    text = _error_text(exc).lower()
    tokens = ["sharing violation", "being used", "in use", "使用中", "ロック", "locked"]
    return any(token in text for token in tokens)


def _emit_word_diagnostics(word) -> None:
    global WORD_DIAGNOSTICS_EMITTED
    with word_diag_lock:
        if WORD_DIAGNOSTICS_EMITTED:
            return
        WORD_DIAGNOSTICS_EMITTED = True
    try:
        version = getattr(word, "Version", "unknown")
    except Exception:
        version = "unknown"
    bitness = "unknown"
    try:
        os_info = str(word.System.OperatingSystem)
        if "64" in os_info:
            bitness = "64-bit"
        elif "32" in os_info:
            bitness = "32-bit"
    except Exception:
        pass
    python_bits = platform.architecture()[0]
    try:
        pywin32_version = getattr(win32com.client, "__version__", "unknown")
    except Exception:
        pywin32_version = "unknown"
    diag = (
        f"WordDiag: word_version={version}; word_bitness={bitness}; "
        f"python_bitness={python_bits}; pywin32={pywin32_version}"
    )
    sys.stderr.write(diag + "\n")
    sys.stderr.flush()


def configure_doc_diag(
    enabled: bool,
    perfile: int,
    exts_argument: str,
    legacy_mode: str,
    single_thread: bool,
) -> None:
    global DOC_DIAG_CONTEXT, DOC_DIAG_EMITTED
    exts_display = (exts_argument or "").strip()
    if not exts_display:
        exts_display = "*"
    DOC_DIAG_CONTEXT = {
        "enabled": bool(enabled),
        "perfile": perfile,
        "exts": exts_display,
        "legacy_mode": (legacy_mode or "com").lower(),
        "single_thread": bool(single_thread),
    }
    DOC_DIAG_EMITTED = not enabled


def _emit_doc_diag_if_needed(word) -> None:
    global DOC_DIAG_EMITTED
    if not DOC_DIAG_CONTEXT.get("enabled"):
        return
    with DOC_DIAG_LOCK:
        if DOC_DIAG_EMITTED:
            return
        DOC_DIAG_EMITTED = True

    arch = platform.architecture()[0]
    if "64" in arch:
        py_bits = "64"
    elif "32" in arch:
        py_bits = "32"
    else:
        py_bits = arch

    word_bits = "unknown"
    try:
        os_info = str(word.System.OperatingSystem)
        if "64" in os_info:
            word_bits = "64"
        elif "32" in os_info:
            word_bits = "32"
    except Exception:
        pass

    gencache_display = "unavailable"
    if win32com is not None:
        try:
            cache_path = win32com.client.gencache.GetGeneratePath()
        except Exception as exc:
            gencache_display = f"error:{_clean_detail(_error_text(exc)) or type(exc).__name__}"
        else:
            gencache_display = cache_path or "unknown"
    gencache_display = _clean_detail(gencache_display) or gencache_display

    perfile = DOC_DIAG_CONTEXT.get("perfile", 0)
    exts_display = DOC_DIAG_CONTEXT.get("exts", "*") or "*"
    legacy_mode = DOC_DIAG_CONTEXT.get("legacy_mode", "com")

    single_thread = DOC_DIAG_CONTEXT.get("single_thread", False)
    diag = (
        f"diag: py={py_bits}, word={word_bits}, gencache={gencache_display}, "
        f"perfile={perfile}, exts={exts_display}, "
        f"legacy-doc-mode={legacy_mode}, single-thread={bool(single_thread)}"
    )
    sys.stderr.write(diag + "\n")
    sys.stderr.flush()
def _build_open_sequences(file_name: str) -> List[dict]:
    base = {
        "FileName": file_name,
        "ReadOnly": True,
        "AddToRecentFiles": False,
        "ConfirmConversions": False,
        "Visible": True,
        "Revert": True,
    }
    sequences: List[dict] = []

    def _append(args: dict) -> None:
        sequences.append(args)

    _append(base.copy())

    for encoding in (1200, 65001):
        args = base.copy()
        args["Encoding"] = encoding
        _append(args)

    alt = base.copy()
    alt["Revert"] = False
    _append(alt)

    for encoding in (1200, 65001):
        args = alt.copy()
        args["Encoding"] = encoding
        _append(args)

    return sequences


def _try_open_document(word, candidates: List[str]) -> Tuple[Optional[object], Optional[dict], Optional[Exception]]:
    last_args: Optional[dict] = None
    last_error: Optional[Exception] = None
    for attempt in range(DOC_LOCK_RETRY_COUNT + 1):
        lock_retry_requested = False
        for candidate in candidates:
            for open_args in _build_open_sequences(candidate):
                try:
                    doc = word.Documents.Open(**open_args)
                    return doc, open_args, None
                except Exception as exc:  # pragma: no cover - depends on Word
                    last_error = exc
                    last_args = open_args
                    if _should_retry_lock(exc):
                        lock_retry_requested = True
                        break
            if lock_retry_requested:
                break
        if lock_retry_requested and attempt < DOC_LOCK_RETRY_COUNT:
            time.sleep(DOC_LOCK_RETRY_DELAY)
            continue
        break
    return None, last_args, last_error


def _describe_word_launch_failure(exc: Exception) -> str:
    reason = "Word を起動できません。Word がインストールされていない、または Office と Python のビット数 (32/64) が一致していない可能性があります。"
    text = _error_text(exc).lower()
    hresult = getattr(exc, "hresult", None)
    if hresult in {-2147221005, -2147221164} or "class not registered" in text:
        reason = "Word がインストールされていないか、Office と Python のビット数 (32/64) が一致していません。"
    elif "server execution failed" in text or hresult in {-2146959355}:
        reason = "Word の COM 自動化を開始できません。Office と Python のビット数 (32/64) を確認してください。"
    return reason


def _convert_doc_via_com(path: str) -> List[str]:
    if pythoncom is None or win32com is None:
        raise DocConversionError(
            "Init",
            "pywin32 がインストールされていないため Word COM を利用できません (必要に応じて 'python -m pywin32_postinstall -install' を実行してください)",
        )

    original_path = path
    temp_path: Optional[str] = None
    word = None
    doc = None
    last_open_args: Optional[dict] = None

    try:
        try:
            word = win32com.client.gencache.EnsureDispatch("Word.Application")
        except Exception as exc:  # pragma: no cover - depends on Word availability
            reason = _describe_word_launch_failure(exc)
            detail = ", ".join(part for part in [reason, _format_com_exception(exc)] if part)
            raise DocConversionError("Launch", detail)

        try:
            word.Visible = True
        except Exception as exc:
            warn_once(
                "word_visible",
                f"Word.Visible を設定できません: {_error_text(exc)}{_format_hresult(exc)}",
            )
        try:
            word.DisplayAlerts = 0
        except Exception as exc:
            warn_once("word_alerts", f"Word.DisplayAlerts を設定できません: {_error_text(exc)}{_format_hresult(exc)}")
        else:
            try:
                _ = word.DisplayAlerts
            except Exception as exc:
                if _looks_like_eula_block(exc):
                    warn_once(
                        "word_eula",
                        "Word の初回起動ダイアログ (EULA) が表示されている可能性があります。Word を手動で起動し、ライセンスに同意してください。",
                    )

        try:
            constants = getattr(win32com.client, "constants", None)
            if constants is not None:
                security_value = getattr(constants, "msoAutomationSecurityForceDisable", MSO_AUTOMATION_SECURITY_FORCE_DISABLE)
            else:
                security_value = MSO_AUTOMATION_SECURITY_FORCE_DISABLE
            word.AutomationSecurity = security_value
        except Exception:
            try:
                word.AutomationSecurity = MSO_AUTOMATION_SECURITY_FORCE_DISABLE
            except Exception as exc:
                warn_once(
                    "word_autosec",
                    f"Word.AutomationSecurity を設定できません: {_error_text(exc)}{_format_hresult(exc)}",
                )

        _emit_word_diagnostics(word)
        _emit_doc_diag_if_needed(word)

        normalized_path = os.path.abspath(path)
        candidates: List[str] = []
        seen = set()

        def _add_candidate(candidate: str) -> None:
            lowered = os.path.normcase(candidate)
            if lowered not in seen:
                seen.add(lowered)
                candidates.append(candidate)

        preferred = ensure_extended_path(normalized_path)
        _add_candidate(preferred)
        _add_candidate(normalized_path)
        _add_candidate(original_path)

        doc, last_open_args, open_error = _try_open_document(word, candidates)
        if doc is None:
            detail_parts = []
            if open_error is not None:
                detail = _format_com_exception(open_error)
                if detail:
                    detail_parts.append(detail)
            detail_args = _summarize_open_args(last_open_args)
            if detail_args:
                detail_parts.append(detail_args.strip(", "))
            raise DocConversionError("Open", ", ".join(detail_parts) or "msg=Open failed")

        _log_doc_stage("Open", path)

        fd, temp_raw = tempfile.mkstemp(suffix=".txt")
        os.close(fd)
        temp_path = os.path.abspath(temp_raw)
        save_target = ensure_extended_path(temp_path)

        try:
            try:
                if hasattr(doc, "SaveAs2"):
                    doc.SaveAs2(save_target, FileFormat=WD_FORMAT_UNICODE_TEXT)
                else:
                    doc.SaveAs(save_target, FileFormat=WD_FORMAT_UNICODE_TEXT)
            except TypeError:
                if hasattr(doc, "SaveAs2"):
                    doc.SaveAs2(save_target, WD_FORMAT_UNICODE_TEXT)
                else:
                    doc.SaveAs(save_target, WD_FORMAT_UNICODE_TEXT)
            except Exception as exc:
                raise DocConversionError("SaveAs", _format_com_exception(exc) or _error_text(exc))

            deadline = time.time() + DOC_SAVEAS_MAX_WAIT
            file_size = 0
            while True:
                try:
                    if temp_path and os.path.exists(temp_path):
                        file_size = os.path.getsize(temp_path)
                        if file_size > 0:
                            break
                except OSError:
                    pass
                if time.time() >= deadline:
                    raise DocConversionError("SaveAs", "msg=SaveAs produced no data")
                time.sleep(DOC_SAVEAS_POLL_INTERVAL)

            _log_doc_stage("SaveAs", temp_path)
        finally:
            try:
                doc.Close(False)
            except Exception as exc:
                _log_doc_error("Close", path, _format_com_exception(exc) or _error_text(exc))
            doc = None

        _log_doc_stage("Read", f"{temp_path} size={file_size}")

        lines: List[str] = []
        try:
            with open(temp_path, "r", encoding="utf-16", errors="replace", newline="") as reader:
                for raw_line in reader:
                    normalized = raw_line.replace("\r", "").rstrip("\n")
                    lines.append(normalized)
        except Exception as exc:
            raise DocConversionError("Read", _format_com_exception(exc) or _error_text(exc))
        return lines
    except DocConversionError:
        raise
    except Exception as exc:
        raise DocConversionError("Unexpected", _format_com_exception(exc) or _error_text(exc))
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception as exc:
                _log_doc_error("Close", path, _format_com_exception(exc) or _error_text(exc))
        if word is not None:
            try:
                word.Quit()
            except Exception as exc:
                _log_doc_error("Quit", path, _format_com_exception(exc) or _error_text(exc))
        if temp_path:
            try:
                os.remove(temp_path)
            except Exception as exc:
                try:
                    os.remove(ensure_extended_path(temp_path))
                except Exception as fallback_exc:
                    detail = _format_com_exception(fallback_exc) or _error_text(fallback_exc)
                    if not detail:
                        detail = _error_text(exc)
                    _log_doc_error("Cleanup", temp_path, detail)


def _find_soffice_executable() -> Optional[str]:
    candidates = ["soffice", "soffice.exe"]
    for candidate in candidates:
        path = shutil.which(candidate)
        if path:
            return path
    return None


def _locate_converted_txt(temp_dir: str, source_path: str) -> Optional[str]:
    expected = os.path.join(temp_dir, os.path.splitext(os.path.basename(source_path))[0] + ".txt")
    if os.path.exists(expected):
        return expected
    try:
        for name in os.listdir(temp_dir):
            if name.lower().endswith(".txt"):
                candidate = os.path.join(temp_dir, name)
                if os.path.isfile(candidate):
                    return candidate
    except Exception:
        return None
    return None


def _convert_doc_via_soffice(path: str) -> List[str]:
    soffice = _find_soffice_executable()
    if not soffice:
        raise DocConversionError("soffice-not-found", "soffice.exe が見つからないため LibreOffice 変換を利用できません")

    temp_dir = tempfile.mkdtemp(prefix="fff_soffice_")
    try:
        command = [
            soffice,
            "--headless",
            "--convert-to",
            "txt:Text",
            "--outdir",
            temp_dir,
            os.path.abspath(path),
        ]
        try:
            completed = subprocess.run(
                command,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                encoding="utf-8",
                errors="replace",
            )
        except FileNotFoundError as exc:
            detail = f"Launch failed: {_error_text(exc)}"
            raise DocConversionError("soffice-launch", detail)
        except Exception as exc:
            detail = f"Execution failed: {_error_text(exc)}"
            raise DocConversionError("soffice-run", detail)
        if completed.returncode != 0:
            message = completed.stderr.strip() or completed.stdout.strip()
            detail = f"stderr={message}" if message else None
            raise DocConversionError(f"soffice-exit={completed.returncode}", detail)

        output_path = _locate_converted_txt(temp_dir, path)
        if not output_path:
            raise DocConversionError("soffice-output-missing", "LibreOffice が txt を生成しませんでした")

        try:
            with open(output_path, "r", encoding="utf-8", errors="replace") as reader:
                return reader.read().splitlines()
        except Exception as exc:
            detail = f"Read failed: {_error_text(exc)}{_format_hresult(exc)}"
            raise DocConversionError("soffice-read", detail)
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def _find_antiword_executable() -> Optional[str]:
    candidates = ["antiword", "antiword.exe"]
    for candidate in candidates:
        path = shutil.which(candidate)
        if path:
            return path
    return None


def _convert_doc_via_antiword(path: str) -> List[str]:
    antiword = _find_antiword_executable()
    if not antiword:
        raise DocConversionError("antiword-not-found", "antiword が PATH 上に見つかりません")

    commands = [
        [antiword, "-m", "UTF-8.txt", os.path.abspath(path)],
        [antiword, os.path.abspath(path)],
    ]
    last_failure: Optional[subprocess.CompletedProcess] = None
    for command in commands:
        try:
            completed = subprocess.run(
                command,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                encoding="utf-8",
                errors="replace",
            )
        except FileNotFoundError as exc:
            detail = f"Launch failed: {_error_text(exc)}"
            raise DocConversionError("antiword-launch", detail)
        except Exception as exc:
            detail = f"Execution failed: {_error_text(exc)}"
            raise DocConversionError("antiword-run", detail)
        if completed.returncode == 0:
            output = completed.stdout or ""
            return output.replace("\r\n", "\n").splitlines()
        last_failure = completed

    if last_failure is not None:
        message = last_failure.stderr.strip() or last_failure.stdout.strip()
        detail = f"stderr={message}" if message else None
        raise DocConversionError(f"antiword-exit={last_failure.returncode}", detail)

    raise DocConversionError("antiword", "antiword での変換に失敗しました")


def iter_doc_legacy_lines(path: str, mode: str) -> Iterable[str]:
    converters = []
    normalized_mode = (mode or "").lower()
    if normalized_mode in {"com", "auto"}:
        converters.append(_convert_doc_via_com)
    if normalized_mode in {"auto", "external"}:
        converters.append(_convert_doc_via_soffice)
        converters.append(_convert_doc_via_antiword)

    for converter in converters:
        try:
            if converter is _convert_doc_via_com:
                global DOC_WORKER_POOL
                if DOC_WORKER_POOL is None:
                    configure_doc_workers(DOC_DIAG_CONTEXT.get("single_thread", False))
                lines = DOC_WORKER_POOL.run(converter, path)
            else:
                lines = converter(path)
        except DocConversionError as exc:
            _log_doc_error(exc.stage, path, exc.detail)
            continue
        except Exception as exc:  # pragma: no cover - unexpected path
            _log_doc_error("unexpected", path, str(exc))
            continue

        for line in lines:
            yield line
        return


def scan_doc_legacy(path: str, matcher, perfile: int, legacy_mode: str) -> int:
    hits = 0
    for lineno, line in enumerate(iter_doc_legacy_lines(path, legacy_mode), 1):
        if matcher(line):
            emit_tsv(path, "", lineno, line)
            hits += 1
            if perfile and hits >= perfile:
                break
    _log_doc_stage("Emit", f"{path} hits={hits}")
    return hits


def _excel_column_name(index: int) -> str:
    name = ""
    i = index
    while i >= 0:
        i, remainder = divmod(i, 26)
        name = chr(65 + remainder) + name
        i -= 1
    return name


def scan_xls_legacy(path: str, matcher, perfile: int) -> int:
    if xlrd is None:
        warn_once("xls", "xlrd<=1.2 がインストールされていないため .xls をスキップします")
        return 0

    try:
        workbook = xlrd.open_workbook(path, on_demand=True)
    except Exception as exc:
        warn_once(f"xls:{path}", f".xls 読み込み失敗: {path} ({exc})")
        return 0

    hits = 0
    try:
        for sheet in workbook.sheets():
            for row_idx in range(sheet.nrows):
                for col_idx in range(sheet.ncols):
                    try:
                        value = sheet.cell_value(row_idx, col_idx)
                    except Exception:
                        continue
                    if value in (None, ""):
                        continue
                    text = str(value).strip()
                    if not text:
                        continue
                    entry = f"{sheet.name}!{_excel_column_name(col_idx)}{row_idx + 1}"
                    if matcher(text):
                        emit_tsv(path, entry, row_idx + 1, text)
                        hits += 1
                        if perfile and hits >= perfile:
                            return hits
    finally:
        try:
            workbook.release_resources()
        except Exception:
            pass
    return hits


def scan_file(path: str, matcher, args, exts: set) -> int:
    ext = normalize_ext(os.path.splitext(path)[1])
    perfile = args.perfile
    if ext == "zip":
        if args.zip:
            return scan_zip(path, matcher, exts, perfile)
        return 0
    if not should_target(path, exts):
        return 0
    if ext == "docx" and args.word:
        return scan_docx(path, matcher, perfile)
    if ext == "xlsx" and args.excel:
        return scan_xlsx(path, matcher, perfile)
    if ext == "doc" and args.legacy and args.word:
        return scan_doc_legacy(path, matcher, perfile, args.legacy_doc)
    if ext == "xls" and args.legacy and args.excel:
        return scan_xls_legacy(path, matcher, perfile)
    return scan_text_file(path, matcher, exts, perfile)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--folder", required=True)
    ap.add_argument("--query", required=True)
    ap.add_argument("--regex", action="store_true")
    ap.add_argument("--zip", action="store_true")
    ap.add_argument("--recursive", action="store_true")
    ap.add_argument("--exts", default="")
    ap.add_argument("--perfile", type=int, default=0)
    ap.add_argument("--word", action="store_true")
    ap.add_argument("--excel", action="store_true")
    ap.add_argument("--legacy", action="store_true")
    ap.add_argument("--legacy-doc", choices=["auto", "com", "external"], default="com")
    ap.add_argument("--doc-single-thread", action="store_true")
    ap.add_argument("--max-workers", type=int, default=0)
    ap.add_argument("--exclude-folders", default="")
    ap.add_argument("--diag", action="store_true")
    args = ap.parse_args()

    args.legacy_doc = (args.legacy_doc or "com").lower()

    if args.perfile is None or args.perfile < 0:
        args.perfile = 0

    if args.diag:
        _emit_startup_diag(args.perfile, args.exts, args.legacy_doc, args.doc_single_thread)
    elif args.exts:
        sys.stderr.write(f"diag: exts={args.exts}\n")
        sys.stderr.flush()

    configure_doc_diag(
        args.diag,
        args.perfile,
        args.exts,
        args.legacy_doc,
        args.doc_single_thread,
    )

    if args.legacy_doc in {"com", "auto"}:
        configure_doc_workers(args.doc_single_thread)
    else:
        shutdown_doc_workers()

    ext_filter = set(
        e.strip().lower()
        for e in args.exts.replace(",", ";").split(";")
        if e.strip()
    )

    exclude_filter = set(
        e.strip().lower()
        for e in args.exclude_folders.replace(",", ";").split(";")
        if e.strip()
    )

    try:
        matcher = build_matcher(args.query, args.regex)
    except re.error as exc:
        sys.stderr.write(f"正規表現エラー: {exc}\n")
        sys.stderr.flush()
        return

    files = list(iter_paths(args.folder, args.recursive, exclude_filter))
    emit_status("queued", len(files))

    max_workers = args.max_workers if args.max_workers > 0 else (os.cpu_count() or 4)
    max_workers = max(1, max_workers)
    processed = 0
    total_hits = 0
    start = time.time()

    def worker(p: str):
        emit_status("current", p)
        try:
            return p, scan_file(p, matcher, args, ext_filter)
        except Exception as exc:  # pragma: no cover - unexpected
            warn_once(f"file:{p}", f"処理失敗: {p} ({exc})")
            return p, 0

    try:
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {executor.submit(worker, p): p for p in files}
            for fut in as_completed(futures):
                path = futures[fut]
                try:
                    _, hits = fut.result()
                except Exception as exc:  # pragma: no cover - already handled
                    warn_once(f"future:{path}", f"処理中に例外: {path} ({exc})")
                    hits = 0
                processed += 1
                total_hits += hits
                emit_status("progress", processed, total_hits, path)

        elapsed = time.time() - start
        emit_status("done", processed, total_hits, f"{elapsed:.3f}")
    finally:
        shutdown_doc_workers()
        try:
            sys.stdout.flush()
        except Exception:
            pass
        try:
            sys.stderr.flush()
        except Exception:
            pass


if __name__ == "__main__":
    main()
